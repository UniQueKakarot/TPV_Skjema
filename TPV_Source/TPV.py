import tkinter as tk
from tkinter import ttk
from tkinter import filedialog
from tkinter import messagebox as mBox
from pathlib import Path
import os.path
from datetime import datetime, date
import webbrowser
import logging

from configobj import ConfigObj
import openpyxl as op
from docx import Document


class TPV_Main():

    def __init__(self, master, tabcontroll, name, config_name):
        
        self.config_name = config_name
        self.date = datetime.today()

        frametxt = 'TPV Skjema'

        if os.path.isfile(self.config_name) is True:
            self.config = ConfigObj(self.config_name, encoding='utf8', default_encoding='utf8')
            frametxt = self.config['Diversje']['5']

        self.master = master
        self.tab = tabcontroll

        self.tabvar = ttk.Frame(self.tab)

        self.tab.add(self.tabvar, text=name)
        self.tab.pack(expand=1, fill="both")

        # Instantiating a labelframe to contain the application in
        self.TPV_Body = ttk.LabelFrame(self.tabvar, text=frametxt)
        self.TPV_Body.pack(expand=1)

        menuBar = tk.Menu(self.master)
        self.master.config(menu=menuBar)

        fileMenu = tk.Menu(menuBar, tearoff=0)
        helpMenu = tk.Menu(menuBar, tearoff=0)

        helpMenu.add_command(label='Prosedyre', command=self.procedure)
        helpMenu.add_command(label='Hjelp', command=self._op_wiki)
        menuBar.add_cascade(label='Info', menu=helpMenu)

        fileMenu.add_command(label='Lagre Utført vedlikehold', command=self.maintanance)
        fileMenu.add_command(label='Åpne excel fil', command=self._op_saved)
        menuBar.add_cascade(label='Alternativer', menu=fileMenu)

        #self.logging()
        self._config_gen()
        self.main()

        #self.master.after(1000, self.persistent_colors)
        #self.persistent_colors()

    def logging(self):

        # Create the Logger
        self.logger = logging.getLogger(__name__)
        self.logger.setLevel(logging.INFO)

        # Create the Handler for logging data to a file
        logger_handler = logging.FileHandler('TPV_log.log')
        logger_handler.setLevel(logging.INFO)

        # Create a Formatter for formatting the log messages
        logger_formatter = logging.Formatter('%(name)s - %(levelname)s - %(message)s')

        # Add the Formatter to the Handler
        logger_handler.setFormatter(logger_formatter)

        # Add the Handler to the Logger
        self.logger.addHandler(logger_handler)

    def _config_gen(self):

        if not os.path.isfile(self.config_name):

            self.config = ConfigObj(encoding='utf8', default_encoding='utf8')
            self.config.filename = self.config_name

            self.config['Vedlikeholdspunkt'] = {}
            self.config['Vedlikeholdspunkt']['1'] = 'Put your info here'
            self.config['Vedlikeholdspunkt']['2'] = 'Put your info here'
            self.config['Vedlikeholdspunkt']['3'] = 'Put your info here'
            self.config['Vedlikeholdspunkt']['4'] = 'Put your info here'
            self.config['Vedlikeholdspunkt']['5'] = 'Put your info here'

            self.config['Handling'] = {}
            self.config['Handling']['1'] = 'Put your info here'
            self.config['Handling']['2'] = 'Put your info here'
            self.config['Handling']['3'] = 'Put your info here'
            self.config['Handling']['4'] = 'Put your info here'
            self.config['Handling']['5'] = 'Put your info here'

            self.config['Oljetype'] = {}
            self.config['Oljetype']['1'] = 'Put your info here'
            self.config['Oljetype']['2'] = 'Put your info here'
            self.config['Oljetype']['3'] = 'Put your info here'
            self.config['Oljetype']['4'] = 'Put your info here'
            self.config['Oljetype']['5'] = 'Put your info here'

            self.config['Hyppighet'] = {}
            self.config['Hyppighet']['1'] = 'Put your info here'
            self.config['Hyppighet']['2'] = 'Put your info here'
            self.config['Hyppighet']['3'] = 'Put your info here'
            self.config['Hyppighet']['4'] = 'Put your info here'
            self.config['Hyppighet']['5'] = 'Put your info here'

            self.config['Diversje'] = {}
            self.config['Diversje']['1'] = 'Du kan skrive ekstra info her:'
            self.config['Diversje']['2'] = 'Annet:'
            self.config['Diversje']['3'] = ''
            self.config['Diversje']['4'] = '20'
            self.config['Diversje']['5'] = 'TPV Skjema for Maskin...'

            self.config['Filbehandling'] = {}
            self.config['Filbehandling']['1'] = ''
            self.config['Filbehandling']['2'] = '850x450'
            self.config['Filbehandling']['3'] = ''
            self.config['Filbehandling']['4'] = ''

            self.config['Vedlikeholdsjekk'] = {}
            self.config['Vedlikeholdsjekk']['1'] = '0'
            self.config['Vedlikeholdsjekk']['2'] = ''
            self.config['Vedlikeholdsjekk']['3'] = ''

            self.config['Ukentlig'] = {}
            self.config['Ukentlig']['1'] = '0'

            self.config['Diversje']['3'] = self.date.strftime('%Y')

            self.config.write()

    def main(self):
        
        """Main body of the gui application"""

        # Accessing the config parser and getting the number of keys
        
        keys = self.config.keys()

        first_key = keys[0]

        # Collecting current date and weekday in full name
        today = self.date.strftime('%A')

        # Collecting which number in the month the current day is
        today_number = self.date.strftime('%d')
        today_number = int(today_number)

        # Collecting which number in the year the current day is
        day_of_year = self.date.strftime('%j')
        day_of_year = int(day_of_year)

        # Empty list assigned for holding info on entries in the first key of the config file
        values = []

        # Establish the number of entries in the config file
        for i in self.config[first_key]:
            values.append(i)

        length = len(values)
        
        day = self.day_check()

        key_weekly = 0
        self.weekly = {}

        key_monthly = 0
        self.monthly = {}

        key_quarterly = 0
        self.quarterly = {}

        key_halfyear = 0
        self.halfyear = {}

        intervalls_counter = 1
        test = 1

        row_ved = 1
        row_han = 1
        row_olj = 1
        row_hyp = 1

        # Generating labels on the fly based on how many entries it is in the config file
        for value in self.config['Vedlikeholdspunkt'].values():
            label = ttk.Label(self.TPV_Body, text=value, font=FONT1)
            label.grid(row=row_ved, column=1, sticky=tk.W, padx=15)
            row_ved += 1

        for value in self.config['Handling'].values():
            label = ttk.Label(self.TPV_Body, text=value, font=FONT1)
            label.grid(row=row_han, column=2, sticky=tk.W, padx=15)
            row_han += 1

        for value in self.config['Oljetype'].values():
            label = ttk.Label(self.TPV_Body, text=value, font=FONT1)
            label.grid(row=row_olj, column=3, sticky=tk.W, padx=15)
            row_olj += 1

        for value in self.config['Hyppighet'].values():

            lowcas = value.lower()

            if lowcas == 'daglig':
                label = ttk.Label(self.TPV_Body, text=value, font=FONT1, background='green')
                label.grid(row=row_hyp, column=4, sticky=tk.W, padx=15)
                row_hyp += 1

            elif lowcas == 'ukentlig':

                weekly_flag = self.config['Ukentlig']['{0}'.format(test)]
                test += 1

                if today == 'Friday' or weekly_flag == '1':

                    label = ttk.Label(self.TPV_Body, text=value, font=FONT1, background='yellow')
                    label.grid(row=row_hyp, column=4, sticky=tk.W, padx=15)
                    row_hyp += 1

                    self.weekly['Weekly' + str(key_weekly)] = intervalls_counter
                    key_weekly += 1

                else:
                    label = ttk.Label(self.TPV_Body, text=value, font=FONT1)
                    label.grid(row=row_hyp, column=4, sticky=tk.W, padx=15)
                    row_hyp += 1

                    self.weekly['Weekly' + str(key_weekly)] = intervalls_counter
                    key_weekly += 1

            elif lowcas == 'månedlig':

                if today_number == day:

                    label = ttk.Label(self.TPV_Body, text=value, font=FONT1, background='orange')
                    label.grid(row=row_hyp, column=4, sticky=tk.W, padx=15)
                    row_hyp += 1

                    self.monthly['Monthly' + str(key_monthly)] = intervalls_counter
                    key_monthly += 1

                else:
                    label = ttk.Label(self.TPV_Body, text=value, font=FONT1)
                    label.grid(row=row_hyp, column=4, sticky=tk.W, padx=15)
                    row_hyp += 1

                    self.monthly['Monthly' + str(key_monthly)] = intervalls_counter
                    key_monthly += 1

            elif lowcas == 'kvartalsvis':

                label = ttk.Label(self.TPV_Body, text=value, font=FONT1, background='orange')
                label.grid(row=row_hyp, column=4, sticky=tk.W, padx=15)
                row_hyp += 1

            elif lowcas == 'halvår':

                if day_of_year == 183:

                    label = ttk.Label(self.TPV_Body, text=value, font=FONT1, background='red')
                    label.grid(row=row_hyp, column=4, sticky=tk.W, padx=15)
                    row_hyp += 1

                    self.halfyear['Halfyear' + str(key_halfyear)] = intervalls_counter
                    key_halfyear += 1

                else:
                    label = ttk.Label(self.TPV_Body, text=value, font=FONT1)
                    label.grid(row=row_hyp, column=4, sticky=tk.W, padx=15)
                    row_hyp += 1

                    self.halfyear['Halfyear' + str(key_halfyear)] = intervalls_counter
                    key_halfyear += 1

            else:
                label = ttk.Label(self.TPV_Body, text=value, font=FONT1)
                label.grid(row=row_hyp, column=4, sticky=tk.W, padx=15)
                row_hyp += 1

            intervalls_counter += 1

        self.results = {}
        check_boxes = {}

        for i in range(length):
            self.results['checkvar{0}'.format(i)] = tk.IntVar()
        
        row = 1
        for item in self.results:
            check_boxes['check{0}'.format(i)] = ttk.Checkbutton(self.TPV_Body, variable=self.results[item]).grid(row=row, column=0, sticky=tk.W)
            row += 1

        lbl1 = tk.Label(self.TPV_Body, text='Vedlikeholdspunkt:', font=FONT2)
        lbl1.grid(row=0, column=1, sticky=tk.N)
        lbl2 = tk.Label(self.TPV_Body, text='Handling:', font=FONT2)
        lbl2.grid(row=0, column=2, sticky=tk.N)
        lbl3 = tk.Label(self.TPV_Body, text='Oljetype:', font=FONT2)
        lbl3.grid(row=0, column=3, sticky=tk.N)
        lbl4 = tk.Label(self.TPV_Body, text='Hyppighet:', font=FONT2)
        lbl4.grid(row=0, column=4, sticky=tk.N)

        header = self.config['Diversje']['1']
        lbl16 = tk.Label(self.TPV_Body, text=header)
        lbl16.grid(row=row, column=0, columnspan=2,  pady=3)
        row += 1

        self.txt = tk.Text(self.TPV_Body, height=3, width=25)
        self.txt.grid(row=row, column=0, columnspan=2)
        row += 1

        button = ttk.Button(self.TPV_Body, text='Lagre', command=self.save)
        button.grid(row=row, column=0, columnspan=2, sticky=tk.W, pady=15)

    def save(self):
        
        values = []
        for i in self.results:
            values.append(self.results[i])

        self.checkbox_values = []
        for i in values:
            self.checkbox_values.append(i.get())
            
        #print(self.checkbox_values)

        # Main code for writing out the excel file used to save the data in comes here:

        # Calling the config parser and accessing filepath if present in config file
        f_name = self.config['Filbehandling']['1']

        # Dealing with dates related to where in the excel file stuff will be saved
        month = str(self.date.strftime('%B'))
        day_as_string = str(int(self.date.strftime('%d')) + 1)
        day_as_int = int(self.date.strftime('%d')) + 1
        today = self.date.strftime('%d.%m.%Y, %a')

        cell = 'A' + day_as_string

        months = ['January', 'February', 'March', 'April', 'May', 'June', 'July',
                  'August', 'September', 'October', 'November', 'December']

        index = [i for i in self.config['Vedlikeholdspunkt'].values()]

        index.append(self.config['Diversje']['2'])


        # Checking if the file exist or not
        if not os.path.isfile(f_name) or self._year_check() == 1:

            f_new = filedialog.asksaveasfilename(title='Select File',
                                                 filetypes=(("Excel files", ".xlsx"),
                                                            ("All files", "*.*")), 
                                                 defaultextension="*.*")

            # Assigning the filename a place in the config file
            self.config['Filbehandling']['1'] = f_new

            wb = op.Workbook()

            for i in months:
                wb.create_sheet(i)

            del wb['Sheet']
            sh = wb.sheetnames

            col = 2

            for m in sh:
                ws = wb[m]
                for i in index:
                    ws.cell(row=1, column=col, value=i)
                    col += 1
                    if col > len(index) + 1:
                        col = 2
                    else:
                        continue

            ws = wb[month]

            ws[cell] = today

            col = 2

            for i in self.checkbox_values:
                ws.cell(row=day_as_int, column=col, value=i)
                col += 1

            try:
                op.writer.excel.save_workbook(wb, f_new)
                mBox.showinfo('', 'Resultater har blitt lagret')
                self.config.write()

            except FileNotFoundError as e:
                self._error_popup('Error!', e)
                #print(e)

        elif os.path.isfile(f_name) and self._year_check() == 0:

            f_name = self.config['Filbehandling']['1']

            wb = op.load_workbook(filename=f_name)
            ws = wb[month]

            ws[cell] = today

            col = 2

            for i in self.checkbox_values:
                ws.cell(row=day_as_int, column=col, value=i)
                col += 1

            textbox = self.txt.get('1.0', tk.END)
            ws.cell(row=day_as_int, column=col, value=textbox)

            self._persistent_weekly(self.checkbox_values, ws)

            op.writer.excel.save_workbook(wb, f_name)

            mBox.showinfo('', 'Resultater har blitt lagret')

        else:
            self._error_popup('Major Error!', 'The whole universe has shattered, take cover!')

    def _year_check(self):

        """A simple method for checking if we have switched year"""

        config_year = self.config['Diversje']['3']

        current_year = self.date.strftime('%Y')

        if config_year != current_year:
            self.config['Diversje']['3'] = current_year
            self.config.write()
            return 1
        else:
            # returning 0 just to indicate that we have indeed not switched year
            return 0

    def _op_saved(self):

        """Lets you open a preexisting excel file"""

        f_exist = filedialog.askopenfilename()

        self.config['Filbehandling']['1'] = f_exist
        self.config.write()

    def _op_wiki(self):

        """Simply opens your default browser and directs you to the wiki"""

        webbrowser.open('https://github.com/UniQueKakarot/TPV_Skjema/wiki')

    def procedure(self):

        """Lets you select a word file that gets linked to in the UI"""

        f_exist = self.config['Filbehandling']['3']

        if os.path.isfile(f_exist) is True:
            os.system(f_exist)

        elif os.path.isfile(f_exist) is False:
            f_exist = filedialog.askopenfilename()
            os.system(f_exist)
            self.config['Filbehandling']['3'] = f_exist
            self.config.write()

        else:
            # insert some logging here when we have figured out if it is needed
            pass
        
    def maintanance(self):
        
        """Adding in a new window for saving extra information on maintenance done"""
        
        self.new_win = tk.Tk()
        self.new_win.title('Skjema')
        self.new_win.geometry('450x280')
        
        today = self.date.strftime('%d.%m.%Y, %a')
        
        form_body = ttk.LabelFrame(self.new_win, text='Skjema for utført Vedlikehold')
        form_body.pack(expand=1)
        
        lbl1 = tk.Label(form_body, text='Dato: ', font=FONT1)
        lbl1.grid(row=0, column=0, sticky=tk.N)
        self.ent1 = ttk.Entry(form_body)
        self.ent1.grid(row=0, column=1, sticky=tk.N)
        self.ent1.insert(0, today)
        
        lbl2 = tk.Label(form_body, text='Hvem utførte vedlikeholdet?: ', font=FONT1)
        lbl2.grid(row=1, column=0, sticky=tk.W, pady=10)
        self.ent2 = ttk.Entry(form_body)
        self.ent2.grid(row=1, column=1, sticky=tk.N, pady=10)
        
        lbl3 = tk.Label(form_body, text='Hva ble gjort:', font=FONT2)
        lbl3.grid(row=2, column=0, sticky=tk.W)
        self.maintanance_txt = tk.Text(form_body, height=3, width=40)
        self.maintanance_txt.grid(row=3, column=0, columnspan=2, sticky=tk.W)
        
        button1 = ttk.Button(form_body, text='Lagre', command=self.maintanance_save)
        button1.grid(row=4, column=0, sticky=tk.N, pady=10)
            
        button2 = ttk.Button(form_body, text='Avslutt', command=self.maintanance_quit)
        button2.grid(row=4, column=1, sticky=tk.N, pady=10)

    def maintanance_save(self):
        
        """Saving method for the main maintainance method"""
        
        f_exist = self.config['Filbehandling']['4']
        
        if os.path.isfile(f_exist) is True:
            
            document = Document(f_exist)
            
            date = self.ent1.get()
            name = self.ent2.get()
            textbox = self.maintanance_txt.get('1.0', tk.END)
            
            document.add_paragraph(date)
            document.add_paragraph(name)
            document.add_paragraph(textbox)
            document.add_paragraph('')
            
            try:
                document.save(f_exist)
                mBox.showinfo('', 'Resultater har blitt lagret')
            except PermissionError as e:
                self._error_popup('Error!', e)
                mBox.showinfo('', 'Resultatene har ikke blitt lagret')
            
        elif os.path.isfile(f_exist) is False:
            
            f_new = filedialog.asksaveasfilename(title='Select File',
                                                 filetypes=(("Word files", ".docx"),
                                                            ("All files", "*.*")),
                                                 defaultextension="*.*")
            
            document = Document()
            
            date = self.ent1.get()
            name = self.ent2.get()
            textbox = self.maintanance_txt.get('1.0', tk.END)
            
            document.add_paragraph(date)
            document.add_paragraph(name)
            document.add_paragraph(textbox)
            document.add_paragraph('')
            
            document.save(f_new)

            self.config['Filbehandling']['4'] = f_new
            self.config.write()
            
            mBox.showinfo('', 'Resultater har blitt lagret')

        else:
            self._error_popup('Error Saving Document', 'Something went wrong when we tried to save the document')

    def maintanance_quit(self):
        
        """Simply just kills the extra save window"""
        
        self.new_win.destroy()

    def day_check(self):

        """Checks to see if the 20th day of the month falls on a weekend"""
        
        saved_day = self.config['Diversje']['4']
        
        year = self.date.strftime('%Y')
        year = int(year)
        month = self.date.strftime('%m')
        month = int(month)
        day = 20
        
        check_day = date(year, month, day).isoweekday()

        if saved_day == '20' and check_day == 6:
            day = day - 1
            self.config['Diversje']['4'] = day
            self.config.write()
            return day
        
        elif saved_day == '20' and check_day == 7:
            day = day + 1
            self.config['Diversje']['4'] = day
            self.config.write()
            return day

        else:
            return 20

    def _persistent_weekly(self, checkbox_values, active_worksheet):

        #today = self.date.strftime('%A')
        today = 'Friday'
        today_number = int(self.date.strftime('%d'))

        #test = int(self.config['Vedlikeholdsjekk']['1'])
        print(checkbox_values)
        config_pos = 1
        if today == 'Friday':
            
            # weekly values need adjusting, list index 1 to high
            for i in self.weekly.values():

                if checkbox_values[i] == 1:
                    self.config['Ukentlig'][str(config_pos)] = '0'
                    config_pos += 1
                    print("Checkbox = 1")

                else:
                    self.config['Ukentlig'][str(config_pos)] = '1'
                    config_pos += 1
                    
            self.config.write()

        # repeat the above and invert the flag writing with the flag as the condition on the outer if


        """
        print(self.weekly)
        print(self.monthly)
        print(self.halfyear)
        """



        pass

    def _error_popup(self, message, issue):

        mBox.showerror('{}'.format(message), '{}'.format(issue))


#win = tk.Tk()
#win.title("TPV Skjema")
#win.geometry("850x460")

#config_name = 'config.ini'
#config_file = Path('TPV-Skjema/config.ini')

FONT1 = ("Calibri", 11)
FONT2 = ("Calibri", 11, "bold", "underline")

#tpv = TPV_Main()

#win.mainloop()
