""" This module is for the extra instance where you can save additional 
    maintanace information in the tpv application """

import tkinter as tk
from tkinter import ttk
from tkinter import filedialog
from tkinter import messagebox as mBox

import os.path

from docx import Document

class Maintanace():
    def __init__(self, config, date, font1, font2):

        self.master = tk.Tk()
        self.config = config
        self.date = date
        self.font1 = font1
        self.font2 = font2

        self.master.title('Skjema')
        self.master.geometry('450x280')
        
        today = self.date.strftime('%d.%m.%Y, %a')
        
        form_body = ttk.LabelFrame(self.master, text='Skjema for utført Vedlikehold')
        form_body.pack(expand=1)
        
        lbl1 = tk.Label(form_body, text='Dato: ', font=self.font1)
        lbl1.grid(row=0, column=0, sticky=tk.N)
        self.ent1 = ttk.Entry(form_body)
        self.ent1.grid(row=0, column=1, sticky=tk.N)
        self.ent1.insert(0, today)
        
        lbl2 = tk.Label(form_body, text='Hvem utførte vedlikeholdet?: ', font=self.font1)
        lbl2.grid(row=1, column=0, sticky=tk.W, pady=10)
        self.ent2 = ttk.Entry(form_body)
        self.ent2.grid(row=1, column=1, sticky=tk.N, pady=10)
        
        lbl3 = tk.Label(form_body, text='Hva ble gjort:', font=self.font2)
        lbl3.grid(row=2, column=0, sticky=tk.W)
        self.maintanance_txt = tk.Text(form_body, height=3, width=40)
        self.maintanance_txt.grid(row=3, column=0, columnspan=2, sticky=tk.W)
        
        button1 = ttk.Button(form_body, text='Lagre', command=self.maintanance_save)
        button1.grid(row=4, column=0, sticky=tk.N, pady=10)
            
        button2 = ttk.Button(form_body, text='Avslutt', command=self.maintanance_quit)
        button2.grid(row=4, column=1, sticky=tk.N, pady=10)

    def maintanance_save(self):

        f_exist = self.config['Filbehandling']['4']
        
        if os.path.isfile(f_exist) is True:
            
            document = Document(f_exist)
            
            date = self.ent1.get()
            name = self.ent2.get()
            textbox = self.maintanance_txt.get('1.0', tk.END)
            
            document.add_paragraph(date)
            document.add_paragraph(name)
            document.add_paragraph(textbox)
            document.add_paragraph('')
            
            try:
                document.save(f_exist)
                mBox.showinfo('', 'Resultater har blitt lagret')
            except PermissionError as e:
                self._error_popup('Error!', e)
                mBox.showinfo('', 'Resultatene har ikke blitt lagret')
            
        elif os.path.isfile(f_exist) is False:
            
            f_new = filedialog.asksaveasfilename(title='Select File',
                                                 filetypes=(("Word files", ".docx"),
                                                            ("All files", "*.*")),
                                                 defaultextension="*.*")
            
            document = Document()
            
            date = self.ent1.get()
            name = self.ent2.get()
            textbox = self.maintanance_txt.get('1.0', tk.END)
            
            document.add_paragraph(date)
            document.add_paragraph(name)
            document.add_paragraph(textbox)
            document.add_paragraph('')
            
            document.save(f_new)

            self.config['Filbehandling']['4'] = f_new
            self.config.write()
            
            mBox.showinfo('', 'Resultater har blitt lagret')

        else:
            self._error_popup('Error Saving Document', 'Something went wrong when we tried to save the document')

    def maintanance_quit(self):
    
        """Simply just kills the extra save window"""
    
        self.master.destroy()

    def _error_popup(self, message, issue):

        mBox.showerror('{}'.format(message), '{}'.format(issue))

        